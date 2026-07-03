import os
from docx import Document
from openpyxl import load_workbook
import fitz  # PyMuPDF
from translator import translate_text


def translate_docx(input_path, output_path, source_language, target_language):
    doc = Document(input_path)

    # Translate normal paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text and run.text.strip():
                run.text = translate_text(
                    run.text,
                    source_language,
                    target_language
                )

    # Translate tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text and run.text.strip():
                            run.text = translate_text(
                                run.text,
                                source_language,
                                target_language
                            )

    doc.save(output_path)


def translate_excel(input_path, output_path, source_language, target_language):
    workbook = load_workbook(input_path)

    for sheet in workbook.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.strip():
                    cell.value = translate_text(
                        cell.value,
                        source_language,
                        target_language
                    )

    workbook.save(output_path)


def translate_pdf(input_path, output_path, source_language, target_language):
    pdf = fitz.open(input_path)
    new_pdf = fitz.open()

    for page in pdf:
        text = page.get_text("text")

        if text.strip():
            translated_text = translate_text(
                text,
                source_language,
                target_language
            )
        else:
            translated_text = ""

        new_page = new_pdf.new_page(
            width=page.rect.width,
            height=page.rect.height
        )

        new_page.insert_textbox(
            fitz.Rect(50, 50, page.rect.width - 50, page.rect.height - 50),
            translated_text,
            fontsize=11,
            fontname="helv"
        )

    new_pdf.save(output_path)
    new_pdf.close()
    pdf.close()


def translate_file(input_path, output_path, source_language, target_language):
    ext = os.path.splitext(input_path)[1].lower()

    if ext == ".docx":
        translate_docx(input_path, output_path, source_language, target_language)

    elif ext == ".xlsx":
        translate_excel(input_path, output_path, source_language, target_language)

    elif ext == ".pdf":
        translate_pdf(input_path, output_path, source_language, target_language)

    else:
        raise ValueError("Unsupported file type. Please upload DOCX, XLSX, or PDF.")